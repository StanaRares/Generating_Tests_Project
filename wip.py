import streamlit as st
from langchain_openai import ChatOpenAI
from langchain_community.document_loaders import PyPDFLoader
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_chroma import Chroma
from langchain.chains import create_retrieval_chain
from langchain.embeddings import HuggingFaceEmbeddings

from langchain.chains.conversation.memory import ConversationBufferMemory
from langchain.chains import LLMChain
from langchain.prompts.chat import ChatPromptTemplate, SystemMessagePromptTemplate, HumanMessagePromptTemplate

import time
from docx import Document
import docx


if 'system_prompt' not in st.session_state:
    st.session_state.system_prompt = '''You are a knowledgeable assistant specialized in creating multiple-response questions for medical students studying surgery. 
    Use the provided medical information to generate a question in Romanian. 
    Ensure that the question is unique, avoid repetitive formats, and adhere strictly to the specified formatting guidelines and question requirements.
    
    **Question Requirements**:
    - Generate exactly 1 question.
    - Each question must have 10 answer options: the first 5 must be correct (Adevarat) and the other 5 must be incorrect (Fals), don't mix them.
    - Avoid questions like: 'what is the definition of...' , 'what is the most frequent...', 'what is the main... ', avoid questions that cannot have 5 correct answers, avoid repeating questions, ensure diversity in topics covered.
    - Ensure the question is written in correct grammatical Romanian language.
    
    **Formatting**:
    [Numar curent]\t[Întrebare, font style: bold]
    a.\t[Varianta de raspuns 1, validation: correct]\t[Adevarat]
    b.\t[Varianta de raspuns 2, validation: correct]\t[Adevarat]
    c.\t[Varianta de raspuns 3, validation: correct]\t[Adevarat]
    d.\t[Varianta de raspuns 4, validation: correct]\t[Adevarat]
    e.\t[Varianta de raspuns 5, validation: correct]\t[Adevarat]
    f.\t[Varianta de raspuns 6, validation: incorrect]\t[Fals]
    g.\t[Varianta de raspuns 7, validation: incorrect]\t[Fals]
    h.\t[Varianta de raspuns 8, validation: incorrect]\t[Fals]
    i.\t[Varianta de raspuns 9, validation: incorrect]\t[Fals]
    j.\t[Varianta de raspuns 10, validation: incorrect]\t[Fals]

    Ensure the output strictly follows this format without additional explanations or deviations.

    Utilize the following context to create a question:
    {context}

    {chat_history}
    '''
    
# Define the local model configuration
llm = ChatOpenAI(
    model="meta-llama/Meta-Llama-3-70B-Instruct",
    openai_api_key="EMPTY",
    openai_api_base="http://localhost:8000/v1",
    temperature=1,
)

# Helper functions
def set_state(i):
    st.session_state.stage = i

def set_custom_state(i):
    st.session_state.stage = i
    start_time = time.time()
    st.session_state.start_time = start_time

def initialize_retriever():
    
    #document loading
    loader = PyPDFLoader("AN IV_Cap 1_Sangerari_txt.pdf")
    documents = loader.load()
    
    #text splitting
    text_splitter = RecursiveCharacterTextSplitter(chunk_size=900, chunk_overlap=200)
    splits = text_splitter.split_documents(documents)
    
    #embedding
    emb = HuggingFaceEmbeddings()
    vectorstore = Chroma.from_documents(documents=splits, embedding=emb)
    
    # Retrieve and generate using the relevant snippets of the document.
    retriever = vectorstore.as_retriever()
    return retriever


def get_chatbot():
    llm = ChatOpenAI(
        model="meta-llama/Meta-Llama-3-70B-Instruct",
        openai_api_key="EMPTY",
        openai_api_base="http://localhost:8000/v1",
        temperature=1,
    )
    
    
    conversational_memory=ConversationBufferMemory(memory_key="chat_history",return_messages=True,input_key="input")
    
    
    system_message_template = st.session_state.system_prompt
    
    system_message_prompt = SystemMessagePromptTemplate.from_template(system_message_template)
    human_template="{input}"
    human_message_prompt = HumanMessagePromptTemplate.from_template(human_template)
    llmchain = LLMChain(llm=llm, prompt=ChatPromptTemplate.from_messages([system_message_prompt,human_message_prompt]), memory = conversational_memory)
    
    
    retriever = initialize_retriever()
    
    return create_retrieval_chain(retriever, llmchain)


# Initialize session state
if 'stage' not in st.session_state:
    st.session_state.stage = 0

if 'skip_generating' not in st.session_state:
    st.session_state.skip_generating = False

if 'quiz' not in st.session_state:
    st.session_state.quiz = ''

if 'quotes' not in st.session_state:
    st.session_state.quotes = ''

if 'start_time' not in st.session_state:
    st.session_state.start_time = 0

if 'end_time' not in st.session_state:
    st.session_state.end_time = 0

# Streamlit app layout and logic
if st.session_state.stage == 0:
    st.title("📝 Turn your notes into a quiz!")
    st.button("Get Started!", use_container_width=True, type='primary', on_click=set_state, args=[1])

elif st.session_state.stage == 1:
    st.subheader("1. Upload your file(s)", divider='red')

    uploaded_file = st.file_uploader("Choose a file", type=['pdf', 'txt', 'docx'])

    if uploaded_file:
        if uploaded_file.type == "application/pdf":
            with open(uploaded_file.name, "wb") as f:
                f.write(uploaded_file.getbuffer())
            loader = PyPDFLoader(uploaded_file.name)
            docs = loader.load()
        else:
            st.warning("Currently, only PDF files are supported.")

        st.session_state.uploaded_files = docs

        st.button("Continue", use_container_width=True, type='primary', on_click=set_custom_state, args=[2])
    else:
        st.warning("Please upload a file!")

elif st.session_state.stage == 2:
    st.subheader('2. Modify the prompt', divider='red')
    st.session_state.system_prompt = st.text_area('Here is the current prompt:', st.session_state.system_prompt, 500)
    st.button("Continue", use_container_width=True, type='primary', on_click=set_custom_state, args=[3])


elif st.session_state.stage == 3:
    st.subheader("3. Time for the quiz!", divider='red')
    if st.session_state.skip_generating == False:
        st.session_state.skip_generating =True
        agent=get_chatbot()
        quiz = ''
        questions = []
        quotes = ''
        for i in range(1,11):
            generated_question = agent.invoke({"input": "Generate a question based only on the context provided."})
            question = generated_question['answer']['text']
            questions.append(question)
            quiz += question + '\n'

        for i in range(1,11):
            generated_quote=agent.invoke({"input": 'Quote me the context that gives the answer to this question:'+ questions[i-1]})
            quote = generated_quote['answer']['text']
            quotes += quote + '\n'

            st.session_state.quiz = quiz
            st.session_state.quotes = quotes

    st.write(st.session_state.quiz)
    st.write()
    st.write(st.session_state.quotes)

    modify_quiz = st.text_input("Modify the quiz:")
    if modify_quiz:
        st.session_state.system_prompt = 'You must take this first in consideration when generating the question: ' + modify_quiz + ' ' + st.session_state.system_prompt
        st.session_state.skip_generating = False
        regenerate_quiz = st.button('Regenerate the questions')
        if regenerate_quiz:
            generate_quiz = agent.invoke({"input": "Use the former quiz and context provided and adhere to the letter this feedback:" + modify_quiz})

    doc = docx.Document()
    doc.add_heading('Quiz', 0)
    for line in st.session_state.quiz.split('\n'):
        doc.add_paragraph(line)
    temp_file = "quiz.docx"
    doc.save(temp_file)
    with open(temp_file, "rb") as file:
        btn = st.download_button(
            label="Download the quiz",
            data=file,
            file_name="quiz.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
    end_time = time.time()
    st.session_state.end_time = end_time
    time_took = end_time - st.session_state.start_time
    st.write("It took ", time_took, " seconds to generate this quiz.")

