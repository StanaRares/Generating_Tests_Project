import streamlit as st
from langchain_openai import ChatOpenAI
from langchain.chains import create_retrieval_chain
from langchain.chains.combine_documents import create_stuff_documents_chain
from langchain_core.prompts import ChatPromptTemplate
from langchain_community.document_loaders import PyPDFLoader
from langchain.vectorstores import Chroma
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain.embeddings import HuggingFaceEmbeddings

import PyPDF2
import time
from docx import Document
import docx


# Define the local model configuration
llm = ChatOpenAI(
    model="meta-llama/Meta-Llama-3-70B-Instruct",
    openai_api_key="EMPTY",
    openai_api_base="http://localhost:8000/v1",
    temperature=1,
)

# Helper functions
def set_state(i):
    st.session_state.stage = i

def set_custom_state(i):
    st.session_state.stage = i
    start_time = time.time()
    st.session_state.start_time = start_time

def trim_response_at_phrase(text):
    return text.split('END OF QUIZ')[0]

def extract_text_from_docx(docx_path):
    doc = Document(docx_path)
    full_text = []
    for para in doc.paragraphs:
        full_text.append(para.text)
    return ' '.join(full_text)

# Initialize session state
if 'stage' not in st.session_state:
    st.session_state.stage = 0

if 'skip_generating' not in st.session_state:
    st.session_state.skip_generating = False

if 'start_time' not in st.session_state:
    st.session_state.start_time = 0

if 'end_time' not in st.session_state:
    st.session_state.end_time = 0

# Streamlit app layout and logic
if st.session_state.stage == 0:
    st.title("📝 Turn your notes into a quiz!")
    skip_generating = st.toggle('Test sidebar')
    if skip_generating:
        st.session_state.skip_generating = skip_generating
        st.write(st.session_state.skip_generating)
        st.button("Get Started!", use_container_width=True, type='primary', on_click=set_state, args=[4])
    else:
        st.button("Get Started!", use_container_width=True, type='primary', on_click=set_state, args=[1])

elif st.session_state.stage == 1:
    st.subheader("1. Upload your file(s)", divider='red')

    uploaded_file = st.file_uploader("Choose a file", type=['pdf', 'txt', 'docx'])

    if uploaded_file:
        if uploaded_file.type == "application/pdf":
            with open(uploaded_file.name, "wb") as f:
                f.write(uploaded_file.getbuffer())
            loader = PyPDFLoader(uploaded_file.name)
            docs = loader.load()
        else:
            st.warning("Currently, only PDF files are supported.")

        st.session_state.uploaded_files = docs

        st.button("Continue", use_container_width=True, type='primary', on_click=set_custom_state, args=[3])
    else:
        st.warning("Please upload a file!")

elif st.session_state.stage == 2:
    st.subheader("2. Upload a template", divider='red')

    model_file = st.file_uploader("Choose a model/template file", type=['pdf', 'txt', 'docx'])

    if model_file:
        if model_file.type == "application/pdf":
            with open(model_file.name, "wb") as f:
                f.write(model_file.getbuffer())
            model_loader = PyPDFLoader(model_file.name)
            model_docs = model_loader.load()
        else:
            st.warning("Currently, only PDF files are supported for the model.")

        st.session_state.model_docs = model_docs

        st.button("Continue", use_container_width=True, type='primary', on_click=set_custom_state, args=[3])
    else:
        st.warning("Please upload a file!")

elif st.session_state.stage == 3:
    st.subheader("3. Time for the quiz!", divider='red')

    text_splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=200)
    splits = text_splitter.split_documents(st.session_state.uploaded_files)

    # model_content = "\n".join([doc.page_content for doc in st.session_state.model_docs])

    vectorstore = Chroma.from_documents(documents=splits, embedding=HuggingFaceEmbeddings())
    retriever = vectorstore.as_retriever()

    system_prompt = """
        You are a knowledgeable assistant specialized in creating multiple-choice quizzes for medical students studying surgery. 
        Use the provided medical information to generate quiz questions in Romanian. 
        Ensure that questions are unique, avoid repetitive formats, and adhere strictly to the specified formatting guidelines.
        
        Utilize the following context to create a quiz:

        {context}

        **Quiz Requirements**:
        - Generate exactly 10 multiple-choice questions.
        - Each question should have 10 answer options: 5 correct (Adevarat) and 5 incorrect (Fals).
        - Avoid repeating question formats and ensure diversity in topics covered.
        - Provide the source and page number for each question.

        **Formatting**:
        [Numar curent]\t[Întrebare, font style: bold]
        a\t[Varianta de raspuns 1, validation: correct]\tAdevarat
        b\t[Varianta de raspuns 2, validation: correct]\tAdevarat
        c\t[Varianta de raspuns 3, validation: correct]\tAdevarat
        d\t[Varianta de raspuns 4, validation: correct]\tAdevarat
        e\t[Varianta de raspuns 5, validation: correct]\tAdevarat
        f\t[Varianta de raspuns 6, validation: incorrect]\tFals
        g\t[Varianta de raspuns 7, validation: incorrect]\tFals
        h\t[Varianta de raspuns 8, validation: incorrect]\tFals
        i\t[Varianta de raspuns 9, validation: incorrect]\tFals
        j\t[Varianta de raspuns 10, validation: incorrect]\tFals
        [Source of the question, include page number]

        Ensure the output strictly follows this format without additional explanations or deviations.
        """

    prompt = ChatPromptTemplate.from_messages(
        [
            ("system", system_prompt),
            ("human", "{input}"),
        ]
    )

    question_answer_chain = create_stuff_documents_chain(llm, prompt)
    rag_chain = create_retrieval_chain(retriever, question_answer_chain)
    results = rag_chain.invoke({"input": "Generate the quiz"})
    final_quiz = results['answer']
    st.write(final_quiz)

    doc = docx.Document()
    doc.add_heading('Quiz', 0)
    for line in final_quiz.split('\n'):
        doc.add_paragraph(line)
    temp_file = "quiz.docx"
    doc.save(temp_file)
    with open(temp_file, "rb") as file:
        btn = st.download_button(
            label="Download the quiz",
            data=file,
            file_name="quiz.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
    end_time = time.time()
    st.session_state.end_time = end_time
    time_took = end_time - st.session_state.start_time
    st.write("It took ", time_took, " seconds to generate this quiz.")
    # Modify content

